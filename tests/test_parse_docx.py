"""Test cho offline/parse_docx.py — kiểm tra toàn vẹn cấu trúc và **đối chiếu
nguyên văn** với paragraph gốc trong docx (không bịa).
"""

from pathlib import Path

import pytest
from docx import Document

from offline import parse_docx
from src import ids
from src.legal_metadata import LawMetadata

DOCX_PATH = Path("data/graph/raw/Luật-41-2024-QH15.docx")


@pytest.fixture(scope="module")
def parsed() -> dict:
    return parse_docx.parse_docx(DOCX_PATH)


@pytest.fixture(scope="module")
def source_paragraphs() -> list[str]:
    """Tất cả paragraph non-empty trong docx, đã strip."""
    doc = Document(str(DOCX_PATH))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


# ---------- Cấu trúc tổng thể ----------


def test_du_11_chuong(parsed):
    assert len(parsed["chapters"]) == 11


def test_du_13_muc(parsed):
    total = sum(len(ch["sections"]) for ch in parsed["chapters"])
    assert total == 13


def test_du_141_dieu(parsed):
    total = sum(len(ch["articles"]) for ch in parsed["chapters"])
    assert total == 141


def test_article_numbers_lien_tiep(parsed):
    nums = [a["number"] for ch in parsed["chapters"] for a in ch["articles"]]
    assert nums == list(range(1, 142))


def test_validate_pass(parsed):
    """Hàm validate() nội bộ phải chạy không lỗi."""
    parse_docx.validate(parsed)


# ---------- ID convention ----------


def test_id_dung_quy_uoc(parsed):
    art_64 = next(a for ch in parsed["chapters"] for a in ch["articles"] if a["number"] == 64)
    assert art_64["id"] == "L41_2024.A64"
    assert art_64["chapter_id"] == "L41_2024.C5"
    assert art_64["section_id"] == "L41_2024.C5.M3"  # Mục 3 - Chế độ hưu trí


def test_id_khong_trung(parsed):
    seen = set()
    for ch in parsed["chapters"]:
        for sec in ch["sections"]:
            assert sec["id"] not in seen
            seen.add(sec["id"])
        for art in ch["articles"]:
            assert art["id"] not in seen
            seen.add(art["id"])
            for cl in art["clauses"]:
                assert cl["id"] not in seen
                seen.add(cl["id"])
                for pt in cl["points"]:
                    assert pt["id"] not in seen
                    seen.add(pt["id"])


# ---------- Metadata ----------


def test_metadata_dung(parsed):
    law = parsed["law"]
    assert law["code"] == "L41_2024"
    assert law["id"] == "L41_2024"
    assert law["full_id"] == "41/2024/QH15"
    assert law["title"] == "Luật Bảo hiểm xã hội"
    assert law["issuer"] == "Quốc hội"
    assert law["issued_date"] == "2024-06-29"
    assert law["effective_date"] == "2025-07-01"
    assert "XV" in law["session"]


# ---------- Đối chiếu NGUYÊN VĂN với docx (không bịa) ----------


def test_dieu_3_khoan_1_dung_nguyen_van(parsed, source_paragraphs):
    """Khoản 1 Điều 3 (định nghĩa BHXH) phải khớp paragraph gốc."""
    art_3 = next(a for ch in parsed["chapters"] for a in ch["articles"] if a["number"] == 3)
    k1 = next(c for c in art_3["clauses"] if c["number"] == 1)
    # Paragraph gốc dạng "1. <text>"
    expected_src = f"1. {k1['text']}"
    assert (
        expected_src in source_paragraphs
    ), "K1 Điều 3 không khớp byte-for-byte với paragraph trong docx"


def test_dieu_2_khoan_1_diem_a_dung_nguyen_van(parsed, source_paragraphs):
    art_2 = next(a for ch in parsed["chapters"] for a in ch["articles"] if a["number"] == 2)
    k1 = next(c for c in art_2["clauses"] if c["number"] == 1)
    pa = next(p for p in k1["points"] if p["letter"] == "a")
    assert f"a) {pa['text']}" in source_paragraphs


def test_dieu_140_khoan_1_dung_hieu_luc(parsed):
    art_140 = next(a for ch in parsed["chapters"] for a in ch["articles"] if a["number"] == 140)
    k1 = next(c for c in art_140["clauses"] if c["number"] == 1)
    assert "01 tháng 7 năm 2025" in k1["text"]


def test_dieu_141_khoan_15_khong_bi_nhiem_postamble(parsed):
    """Đoạn 'Luật này được Quốc hội ... thông qua...' không được dính vào K15."""
    art_141 = next(a for ch in parsed["chapters"] for a in ch["articles"] if a["number"] == 141)
    k15 = next(c for c in art_141["clauses"] if c["number"] == 15)
    assert "thông qua ngày" not in k15["text"]
    assert "Quốc hội" not in k15["text"] or k15["text"].strip().endswith(
        "Chính phủ quy định chi tiết Điều này."
    )


def test_postamble_co_ratification(parsed):
    assert any("thông qua ngày" in p for p in parsed["postamble"])


# ---------- Tính toàn vẹn nội dung ----------


def test_moi_dieu_co_text_non_empty(parsed):
    for ch in parsed["chapters"]:
        for art in ch["articles"]:
            assert art["text"].strip(), f"Điều {art['number']} text rỗng"


def test_moi_khoan_co_text_non_empty(parsed):
    for ch in parsed["chapters"]:
        for art in ch["articles"]:
            for cl in art["clauses"]:
                assert cl["text"].strip(), f"{cl['id']} text rỗng"


def test_moi_diem_co_text_non_empty(parsed):
    for ch in parsed["chapters"]:
        for art in ch["articles"]:
            for cl in art["clauses"]:
                for pt in cl["points"]:
                    assert pt["text"].strip(), f"{pt['id']} text rỗng"


def test_article_text_chua_title(parsed):
    """Article.text phải có header 'Điều N. Tên điều' ở đầu."""
    for ch in parsed["chapters"]:
        for art in ch["articles"]:
            assert art["text"].startswith(
                f"Điều {art['number']}. "
            ), f"Điều {art['number']}: text không bắt đầu bằng header chuẩn"


# ---------- Reverse provenance (truy ngược ID) ----------


def test_parse_id_round_trip(parsed):
    """Từ ID bất kỳ trong cây, parse_id() trả về đúng (article, clause, point)."""
    art = next(a for ch in parsed["chapters"] for a in ch["articles"] if a["number"] == 64)
    k1 = art["clauses"][0]
    pa = k1["points"][0]

    p_art = ids.parse_id(art["id"])
    assert p_art["article"] == 64

    p_cl = ids.parse_id(k1["id"])
    assert p_cl["article"] == 64
    assert p_cl["clause"] == k1["number"]

    p_pt = ids.parse_id(pa["id"])
    assert p_pt["article"] == 64
    assert p_pt["clause"] == k1["number"]
    assert p_pt["point"] == pa["letter"]


def test_section_chi_co_o_chuong_co_muc(parsed):
    """Article có section_id phải khớp với 1 section trong chương cha."""
    for ch in parsed["chapters"]:
        sec_ids = {s["id"] for s in ch["sections"]}
        for art in ch["articles"]:
            if art["section_id"]:
                assert art["section_id"] in sec_ids, (
                    f"{art['id']} có section_id={art['section_id']} "
                    f"nhưng không thuộc Chương {ch['number']}"
                )
            # Nếu chương có Mục, tất cả Điều phải thuộc một Mục nào đó
            if ch["sections"]:
                assert art["section_id"] is not None, (
                    f"Chương {ch['number']} có {len(ch['sections'])} Mục "
                    f"nhưng Điều {art['number']} không thuộc Mục nào"
                )


# ============================================================================
# allow_no_chapter — văn bản dưới luật (NĐ/QĐ/TT) không có "Chương"
# ============================================================================


def _meta_for(law_id: str, source: Path, *, allow_no_chapter: bool) -> LawMetadata:
    return LawMetadata(
        id=law_id,
        code=law_id,
        full_id=law_id,
        title=f"{law_id} title",
        canonical_title=f"{law_id} canonical",
        type="decree",
        hierarchy_level="nghị định",
        priority=80,
        source_file=source,
        allow_no_chapter=allow_no_chapter,
    )


def _write_paragraphs(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def test_allow_no_chapter_off_default_giu_nguyen_hanh_vi(tmp_path):
    """Mặc định (flag tắt) — Điều không có Chương cha vẫn bị bỏ qua như cũ."""
    src = tmp_path / "no_chapter_default.docx"
    _write_paragraphs(
        src,
        [
            "CHÍNH PHỦ",
            "NGHỊ ĐỊNH",
            "Điều 1. Phạm vi điều chỉnh",
            "1. Nội dung khoản 1 của Điều 1.",
            "Điều 2. Đối tượng áp dụng",
            "1. Nội dung khoản 1 của Điều 2.",
        ],
    )
    meta = _meta_for("L143_2018", src, allow_no_chapter=False)
    result = parse_docx.parse_docx(src, metadata=meta)
    assert result["chapters"] == [], (
        "Khi flag tắt, không được tự tạo Chương — phải giữ hành vi cũ"
    )


def test_allow_no_chapter_on_synth_chuong_va_nhan_dieu(tmp_path):
    """Khi flag bật — synth Chương I và nhận đầy đủ Điều/Khoản."""
    src = tmp_path / "no_chapter_on.docx"
    _write_paragraphs(
        src,
        [
            "CHÍNH PHỦ",
            "NGHỊ ĐỊNH",
            "Điều 1. Phạm vi điều chỉnh",
            "1. Nghị định này quy định phạm vi.",
            "2. Một số trường hợp đặc biệt.",
            "Điều 2. Đối tượng áp dụng",
            "1. Nội dung khoản 1.",
            "a) Điểm a của khoản 1.",
            "b) Điểm b của khoản 1.",
        ],
    )
    meta = _meta_for("L143_2018", src, allow_no_chapter=True)
    result = parse_docx.parse_docx(src, metadata=meta)

    assert len(result["chapters"]) == 1
    ch = result["chapters"][0]
    assert ch["id"] == "L143_2018.C1"
    assert ch["number"] == 1
    assert ch["roman"] == "I"
    assert ch["title"] == meta.canonical_title  # không hardcode chuỗi

    assert [a["number"] for a in ch["articles"]] == [1, 2]
    a1, a2 = ch["articles"]
    assert a1["id"] == "L143_2018.A1"
    assert [c["number"] for c in a1["clauses"]] == [1, 2]
    assert a2["clauses"][0]["points"][0]["letter"] == "a"
    assert a2["clauses"][0]["points"][1]["letter"] == "b"


def test_allow_no_chapter_validate_pass_voi_expected_synth(tmp_path):
    """Validate phải pass khi YAML khai báo `expected.chapters=1` cho synth."""
    src = tmp_path / "no_chapter_validate.docx"
    _write_paragraphs(
        src,
        [
            "Điều 1. Tiêu đề",
            "1. Khoản 1.",
            "Điều 2. Tiêu đề 2",
            "1. Khoản 1.",
        ],
    )
    meta = _meta_for("L146_2018", src, allow_no_chapter=True)
    result = parse_docx.parse_docx(src, metadata=meta)
    # Synth Chương 1, 2 Điều, 0 Mục — validate phải pass với expected khớp.
    parse_docx.validate(
        result, expect_chapters=1, expect_sections=0, expect_articles=2
    )


def test_allow_no_chapter_synth_voi_muc_truoc_dieu(tmp_path):
    """Document mở đầu bằng Mục (không có Chương) — synth Chương vẫn pickup Mục."""
    src = tmp_path / "muc_before_dieu.docx"
    _write_paragraphs(
        src,
        [
            "Mục 1. Quy định chung",
            "Điều 1. Phạm vi",
            "1. Nội dung.",
            "Mục 2. Quy định riêng",
            "Điều 2. Đối tượng",
            "1. Nội dung.",
        ],
    )
    meta = _meta_for("L158_2025", src, allow_no_chapter=True)
    result = parse_docx.parse_docx(src, metadata=meta)

    assert len(result["chapters"]) == 1
    ch = result["chapters"][0]
    assert [s["number"] for s in ch["sections"]] == [1, 2]
    assert [a["number"] for a in ch["articles"]] == [1, 2]
    # Điều 1 thuộc Mục 1, Điều 2 thuộc Mục 2
    assert ch["articles"][0]["section_id"] == "L158_2025.C1.M1"
    assert ch["articles"][1]["section_id"] == "L158_2025.C1.M2"


def test_l41_2024_metadata_default_allow_no_chapter_false():
    """Luật 41/2024 trong YAML không khai báo flag → mặc định False — không đổi."""
    from src.legal_metadata import metadata_for

    meta = metadata_for("L41_2024")
    assert meta.allow_no_chapter is False
